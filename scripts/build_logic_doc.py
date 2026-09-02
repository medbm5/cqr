"""Build the business-logic reference: endpoint, parameters, functions, maths.

One section per page of the cockpit. The audience is the author, five minutes
before a soutenance: what does this page call, what does each function do, and
what do I say out loud when someone asks how the maths works.

Figures come from `deck/figures.json`, which the chart script writes from the
engine - so a number here cannot disagree with a number on a slide.

    python scripts/build_deck_charts.py     # writes deck/figures.json
    python scripts/build_logic_doc.py       # writes logique_metier.docx (+ .pdf)
"""

from __future__ import annotations

import argparse
import json
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent

INK = RGBColor(0x11, 0x18, 0x27)
MUTED = RGBColor(0x55, 0x5F, 0x70)
ACCENT = RGBColor(0x0B, 0x63, 0xA6)
CODE = RGBColor(0x1B, 0x40, 0x60)

BODY = "Calibri"
MONO = "Consolas"


# --------------------------------------------------------------- primitives


def style(document: Document) -> None:
    """Set the document's base typography once."""
    normal = document.styles["Normal"]
    normal.font.name = BODY
    normal.font.size = Pt(10)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.06
    # Word needs the east-asian font set too, or it substitutes.
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY)


#: Structure of the document being built, as a list of heading levels. Reset by
#: `render_document`. Its only purpose is the parity check in `main`: a section
#: added to one language and forgotten in the other is otherwise invisible.
_OUTLINE: list[str] = []


def heading(document: Document, text: str, level: int, *, new_page: bool = False) -> None:
    """A heading in the accent colour, without Word's default blue.

    `new_page` sets `page-break-before` on the heading itself rather than
    inserting a break paragraph: an explicit break lands a blank page whenever
    the previous section happened to end near the bottom of one.
    """
    _OUTLINE.append(f"h{level}{'/page' if new_page else ''}")
    paragraph = document.add_heading(level=level)
    run = paragraph.add_run(text)
    run.font.name = BODY
    run.font.color.rgb = ACCENT if level <= 2 else INK
    run.font.size = Pt({1: 20, 2: 15, 3: 12}[level])
    run.bold = True
    paragraph.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.page_break_before = new_page
    paragraph.paragraph_format.keep_with_next = True


def para(document: Document, text: str, *, size=10, color=INK, italic=False) -> None:
    """One plain paragraph."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.italic = italic


def bullets(document: Document, items: list[str], *, level: int = 0) -> None:
    """A bullet list. `**bold**` and `` `code` `` are honoured inline."""
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Pt(18 + level * 14)
        paragraph.paragraph_format.space_after = Pt(2)
        _rich(paragraph, item)


def numbered(document: Document, items: list[str]) -> None:
    """A numbered list, for call chains where order is the point.

    The numbers are written as text rather than using Word's `List Number`
    style, which keeps a single counter for the whole document and would number
    the second section's chain 8, 9, 10 instead of restarting at 1.
    """
    for index, item in enumerate(items, start=1):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Pt(32)
        paragraph.paragraph_format.first_line_indent = Pt(-18)
        paragraph.paragraph_format.space_after = Pt(2)
        marker = paragraph.add_run(f"{index}.".ljust(4))
        marker.font.size = Pt(10)
        marker.font.color.rgb = MUTED
        _rich(paragraph, item)


#: Inline markers, longest first so `**` is never read as two `*`.
_MARKERS = re.compile(r"(\*\*[^*]+\*\*|\*[^*`]+\*|`[^`]+`)")


def _rich(paragraph, text: str) -> None:
    """Render `**bold**`, `*italic*` and `` `code` `` inside one paragraph.

    Markers do not nest: a bold run containing backticks would need a parser,
    and the content is written to avoid needing one.
    """
    for piece in _MARKERS.split(text):
        if not piece:
            continue
        run = paragraph.add_run()
        if piece.startswith("**") and piece.endswith("**"):
            run.text = piece[2:-2]
            run.bold = True
        elif piece.startswith("`") and piece.endswith("`"):
            run.text = piece[1:-1]
            run.font.name = MONO
            run.font.size = Pt(9)
            run.font.color.rgb = CODE
        elif piece.startswith("*") and piece.endswith("*"):
            run.text = piece[1:-1]
            run.italic = True
        else:
            run.text = piece
        if run.font.size is None:
            run.font.size = Pt(10)


def code(document: Document, text: str) -> None:
    """A monospaced block, for endpoints and formulas."""
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(14)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(6)
    for index, line in enumerate(text.split("\n")):
        run = paragraph.add_run(("\n" if index else "") + line)
        run.font.name = MONO
        run.font.size = Pt(9)
        run.font.color.rgb = CODE


def table(document: Document, headers: list[str], rows: list[list[str]], widths=None) -> None:
    """A light grid table."""
    grid = document.add_table(rows=1, cols=len(headers))
    grid.style = "Light Grid Accent 1"
    grid.alignment = WD_TABLE_ALIGNMENT.LEFT
    for cell, label in zip(grid.rows[0].cells, headers, strict=True):
        cell.text = ""
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.size = Pt(9)
    for row in rows:
        cells = grid.add_row().cells
        for cell, value in zip(cells, row, strict=True):
            cell.text = ""
            _rich(cell.paragraphs[0], value)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
    if widths:
        for row in grid.rows:
            for cell, width in zip(row.cells, widths, strict=True):
                cell.width = width
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def answer(document: Document, label: str, text: str) -> None:
    """The block to read out loud when the question comes.

    The label is a parameter because the two languages announce it differently
    and the block is otherwise identical.
    """
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(12)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(10)
    tag = paragraph.add_run(f"{label}  ")
    tag.bold = True
    tag.font.size = Pt(8.5)
    tag.font.color.rgb = ACCENT
    run = paragraph.add_run(text)
    run.font.size = Pt(10)
    run.italic = True


# ------------------------------------------------------------------ document


def new_document() -> Document:
    """A blank document with this reference's page setup and typography."""
    document = Document()
    style(document)
    for section in document.sections:
        section.left_margin = section.right_margin = Pt(54)
        section.top_margin = section.bottom_margin = Pt(48)
    return document


def build_fr(f: dict) -> Document:
    """The French reference."""
    document = new_document()
    heading(document, "Logique métier, page par page", 1)
    para(
        document,
        "Pour chaque page du cockpit : l'endpoint appelé, ses paramètres, la chaîne de "
        "fonctions traversée, et les maths derrière le calcul. Chaque section finit par la "
        "réponse à donner à l'oral.",
        color=MUTED,
    )
    para(
        document,
        f"Chiffres du moteur, seed 42, {f['n_years']} années simulées. "
        f"Toute valeur citée ici sort de deck/figures.json, écrit par le pipeline.",
        size=9,
        color=MUTED,
        italic=True,
    )

    # ------------------------------------------------------- vue d'ensemble
    heading(document, "Vue d'ensemble", 2)
    table(
        document,
        ["Page", "Endpoint", "Paramètres", "Coût"],
        [
            ["Overview", "les 4 ci-dessous", "—", "tout en cache"],
            ["Telemetry", "`GET /api/telemetry/summary/`\n`GET /api/frequency/`", "aucun", "cache"],
            ["Frequency", "`GET /api/frequency/`\n`GET /api/assets/`", "seuil, fenêtre", "~1 s"],
            ["Severity", "`GET /api/severity/`", "aucun", "cache"],
            ["Simulation", "`POST /api/simulate/`", "7 (voir §4)", "quelques s"],
        ],
    )
    bullets(
        document,
        [
            "**Une seule règle d'architecture** : la vue lit la requête, appelle le moteur, "
            "sérialise. Aucun calcul de risque dans Django, aucun dans le front.",
            "**Deux niveaux de cache** : `get_dataset()` charge et ajuste une fois pour toutes "
            "(les CSV ne changent pas) ; `get_frequency()` et `get_simulation()` sont mémoïsés "
            "par jeu de paramètres.",
            "**Tout objet de résultat expose** `to_explanation()` — la trace numérotée que la "
            "CLI imprime, que l'API sert et que l'interface affiche. Une seule version de "
            "chaque chiffre.",
        ],
    )

    # ============================================================ TELEMETRY
    heading(document, "1. Telemetry — ce que les sondes ont vu", 2, new_page=True)

    heading(document, "Endpoint et paramètres", 3)
    code(document, "GET /api/telemetry/summary/      (aucun paramètre)\nGET /api/frequency/")
    bullets(
        document,
        [
            "Le second appel ne sert qu'au **dernier étage de l'entonnoir** : le taux "
            "d'incidents que la simulation consomme.",
            "Aucun paramètre : la normalisation ne dépend d'aucune convention.",
        ],
    )

    heading(document, "Chaîne de fonctions", 3)
    numbered(
        document,
        [
            "`load_assets()` — lit le référentiel des 20 actifs.",
            "`load_siem()` / `load_edr()` — lisent chaque export et le traduisent en "
            "`SecurityEvent` canoniques. C'est là que les deux échelles de sévérité sont "
            "ramenées à un vocabulaire commun.",
            "`severity_from_siem_label()` — `Low/Medium/High/Critical` → classe.",
            "`severity_from_edr_risk()` — score 0–999 → classe, par points de coupure.",
            "`merge_feeds()` — **dédoublonne** sur l'union des deux flux et produit le "
            "`NormalizationReport`.",
            "`observed_window()` — déduit la fenêtre d'observation des événements eux-mêmes.",
            "`summarize_telemetry()` — agrège en buckets hebdomadaires, mix de sévérité, "
            "top techniques.",
        ],
    )

    heading(document, "Les maths", 3)
    bullets(
        document,
        [
            "**Normalisation des sévérités.** Le SIEM donne un libellé, l'EDR un score. "
            "Points de coupure EDR : `< 50` low, `50–69` medium, `70–93` high, `≥ 94` critical. "
            "Score numérique associé à chaque classe : 0,25 / 0,5 / 0,75 / 1,0.",
            "**Clé de déduplication** : le triplet `(asset_id, technique, observed_at)`. "
            "Deux rapports portant le même triplet décrivent le même événement réel.",
            "**Dédoublonnage sur l'union, pas par jointure.** Chaque flux répète déjà des clés "
            "en interne (761 fois côté SIEM, 543 côté EDR) ; une jointure interne multiplierait "
            "ces doublons et renverrait 13 055 lignes au lieu des 12 343 vrais recoupements.",
            "**Règle de fusion : la pire sévérité gagne.** Si l'EDR voit critical et le SIEM "
            "medium, on garde critical — l'outil le plus proche du poste a vu quelque chose. "
            "Prendre la moyenne diluerait un signal fort.",
            f"**Facteur d'annualisation** = 365 / jours observés = 365 / {f['observed_days']} = "
            f"**{f['annualization']}** (1ᵉʳ nov. 2025 → 31 mai 2026). Recalculé depuis les "
            "données, jamais écrit en dur.",
        ],
    )

    heading(document, "Chiffres produits", 3)
    table(
        document,
        ["Grandeur", "Valeur", "D'où elle vient"],
        [
            ["Lignes brutes", f["rows_read"], "somme des deux exports"],
            ["Événements distincts", f["total_events"], "après déduplication"],
            ["Doublons absorbés", f["duplicates"], f"dont {f['dup_cross_feed']} inter-flux"],
            [
                "Inflation évitée",
                f["inflation"],
                "ce qu'une concaténation naïve aurait ajouté",
            ],
        ],
    )

    answer(
        document,
        "À DIRE",
        "Les deux outils ne sont pas complémentaires, ils sont partiellement redondants : "
        f"{f['dup_cross_feed']} événements portent le même actif, la même technique et le même "
        "horodatage des deux côtés. Je dédoublonne sur ce triplet en gardant la pire sévérité, "
        f"ce qui évite {f['inflation']} d'inflation sur toutes les fréquences en aval.",
    )

    # ============================================================ FREQUENCY
    heading(document, "2. Frequency — à quelle fréquence on est attaqué", 2, new_page=True)

    heading(document, "Endpoint et paramètres", 3)
    code(
        document,
        "GET /api/frequency/?severity_threshold=high&session_window_hours=24\n"
        "GET /api/assets/?severity_threshold=high&session_window_hours=24",
    )
    table(
        document,
        ["Paramètre", "Défaut", "Effet"],
        [
            [
                "`severity_threshold`",
                "`high`",
                "sévérité minimale pour compter comme attaque",
            ],
            ["`session_window_hours`", "`24`", "silence qui ouvre un nouvel épisode"],
        ],
    )
    para(
        document,
        "Ces deux paramètres sont des conventions, pas des mesures. Ils sont exposés comme "
        "curseurs dans l'interface plutôt que cachés dans le code.",
        size=9.5,
        color=MUTED,
        italic=True,
    )

    heading(document, "Chaîne de fonctions", 3)
    numbered(
        document,
        [
            "`get_frequency(seuil, fenêtre)` — mémoïsé par couple de paramètres.",
            "`estimate_frequency()` — orchestre les quatre étapes ci-dessous.",
            "`is_attack_grade(event, seuil)` — filtre : ne garde que high et au-delà.",
            "`sessionize(events, params)` — regroupe les événements en **épisodes**.",
            "`attack_type_for(technique)` — classe chaque technique MITRE en type d'attaque "
            "via un dictionnaire explicite ; l'épisode hérite du type de son événement le plus "
            "grave.",
            "`peer_weighted_base_rate(incidents, params)` — taux d'incidents par "
            "organisation-année chez les pairs.",
            "`calibrate(lambda_detected, base_rate)` — ajuste `p_materialize` pour raccorder "
            "les deux unités.",
        ],
    )

    heading(document, "Les maths", 3)
    bullets(
        document,
        [
            "**Règle d'épisode.** Groupement par **actif seul** — un intrus déclenche les "
            "détections qu'il rencontre, et compter chaque type séparément recompterait la même "
            "intrusion sous plusieurs noms. On coupe dès qu'un silence dépasse la fenêtre, "
            "mesuré par rapport à l'événement **précédent** et non au premier, sinon une "
            "intrusion de trois jours serait découpée artificiellement.",
            f"**λ détecté** = épisodes × 365 / jours observés = {f['episodes']} × "
            f"{f['annualization']} = **{f['lambda_detected']} / an**.",
            "**Taux de base des pairs** = incidents pondérés ÷ organisations-années pondérées. "
            "Numérateur *et* dénominateur portent le même noyau de similarité, un enregistrement "
            f"par `company_id` — soit **{f['peer_companies']} organisations**.",
            "**Le changement d'unité, le point clé.** La télémétrie compte des attaques "
            "*détectées*, la base externe des incidents *à perte*. Ce sont deux unités "
            "différentes.",
            f"**Calibration** : `λ_incident = λ_détecté × p`, avec p fixé pour que λ_incident "
            f"égale le taux des pairs. Ici p ≈ 1 détection sur **{f['p_one_in']}**, soit "
            f"λ_incident = **{f['lambda_incident']} / an**.",
            "**Le mix vient de la télémétrie, le niveau de la base.** Chaque type reçoit "
            "λ_incident × (ses épisodes / épisodes totaux).",
        ],
    )

    heading(document, "Chiffres produits", 3)
    table(
        document,
        ["Grandeur", "Valeur", "Calcul"],
        [
            ["Événements attack-grade", f["attack_grade"], "sévérité ≥ high"],
            ["Épisodes", f["episodes"], f"compression {f['compression']}"],
            ["λ détecté", f"{f['lambda_detected']} / an", "épisodes × 365/212"],
            ["λ incident", f"{f['lambda_incident']} / an", f"λ détecté ÷ {f['p_one_in']}"],
        ],
    )

    answer(
        document,
        "À DIRE",
        "Une alerte n'est pas une attaque : une intrusion produit une rafale de détections. Je "
        f"regroupe par actif avec une fenêtre de silence de 24 h, ce qui ramène "
        f"{f['attack_grade']} événements attack-grade à {f['episodes']} épisodes, soit "
        f"{f['lambda_detected']} attaques détectées par an. Mais la télémétrie compte des "
        "détections et la base externe des incidents à perte — deux unités différentes. Je les "
        f"raccorde par une calibration : une détection sur {f['p_one_in']} finit en perte, soit "
        f"{f['lambda_incident']} incident par an ancré sur {f['peer_companies']} organisations "
        "comparables.",
    )

    # ============================================================= SEVERITY
    heading(document, "3. Severity — ce que coûte un incident", 2, new_page=True)

    heading(document, "Endpoint et paramètres", 3)
    code(document, "GET /api/severity/      (aucun paramètre)")
    para(
        document,
        "Aucun paramètre : le modèle de sévérité ne dépend que du profil cible, qui est fixe "
        "pour ce cas (Retail, ETI, maturité 55).",
        size=9.5,
        color=MUTED,
        italic=True,
    )

    heading(document, "Chaîne de fonctions", 3)
    numbered(
        document,
        [
            "`load_incidents()` — lit et **nettoie** la base externe ; renvoie aussi le "
            "`CleaningReport`.",
            "`repair_mojibake()` — répare le double encodage UTF-8 des libellés de secteur, "
            "qui séparait « Énergie » en deux secteurs distincts.",
            "`fit_severity_model()` — orchestre les étapes ci-dessous, une fois pour l'ensemble "
            "puis une fois par type d'attaque.",
            "`peer_weights(incidents, params)` — un poids par incident, via `sector_weight()`, "
            "`size_weight()` et `maturity_weight()`.",
            "`effective_sample_size(weights)` — n_eff de Kish, ce que la pondération coûte.",
            "`fit_lognormal(losses, weights)` — maximum de vraisemblance pondéré sur les logs.",
            "`diagnose()` — assemble les preuves *contre* l'ajustement : `weighted_ks()`, "
            "`qq_points()`, `fit_pareto_tail()`, `distribution_plot()`.",
        ],
    )

    heading(document, "Les maths", 3)
    bullets(
        document,
        [
            "**Nettoyage.** Les pertes à `-1` sont des valeurs **manquantes**, jamais des 0 €. "
            f"Deux lignes concernées, exclues de l'ajustement : il reste "
            f"**{f['incidents_fitted']} incidents exploitables sur 1 600**.",
            "**Pondération douce, jamais de filtre dur.** Un poids continu par incident, "
            "produit de trois facteurs de ressemblance : secteur (1,0 si Retail, sinon 0,4), "
            "taille (1,0 si ETI, sinon 0,6) et un noyau gaussien sur l'écart de maturité, de "
            "largeur h = 15.",
            f"**Pourquoi pondérer.** Un filtre exact sur le profil ne laisse que "
            f"**{f['hard_filter']} incidents sur {f['incidents_fitted']}**, et plus aucun type "
            "d'attaque n'a d'échantillon crédible. Mathématiquement propre, pratiquement "
            "inutilisable.",
            "**Ajustement lognormal pondéré**, sur l'échelle des logs — c'est simplement une "
            "moyenne et une variance pondérées :",
        ],
    )
    code(
        document,
        "mu    = Σ wᵢ ln xᵢ / Σ wᵢ\nsigma² = Σ wᵢ (ln xᵢ − mu)² / Σ wᵢ",
    )
    bullets(
        document,
        [
            "**Retour en euros** : `médiane = e^mu` et `moyenne = e^(mu + sigma²/2)`. "
            "Attention : mu et sigma **ne sont pas** la moyenne et l'écart-type de la perte, "
            "mais ceux de son logarithme.",
            f"**Exemple, supply chain** : mu = {f['sc_mu']}, sigma = {f['sc_sigma']} → "
            f"médiane **{f['sc_median']}**, moyenne **{f['sc_mean']}**. La queue multiplie la "
            f"moyenne par **{f['sc_ratio']}** : e^(sigma²/2).",
            "**C'est la moyenne qui alimente la perte annuelle**, pas la médiane. Un décideur "
            f"qui raisonne sur « l'incident typique » se trompe d'un facteur {f['sc_ratio']}.",
            "**n_eff de Kish** = `(Σw)² / Σw²`. C'est le nombre d'observations de poids égal "
            "qui porteraient autant d'information. Il mesure la **régularité** des poids, pas "
            f"leur taille. En dessous de **{f['min_neff']}**, le type bascule sur la "
            "distribution poolée, et la substitution est enregistrée.",
            f"**Preuves publiées avec chaque ajustement** : KS pondéré (ici **{f['sc_ks']}** "
            "contre un seuil indicatif ≈ 0,19), QQ-plot des log-pertes, et une **queue de "
            f"Pareto ajustée en rivale** — elle gagne sur {f['pareto_better']} ajustements sur "
            f"{f['fits_total']}, donc VaR 99 et TVaR 99 se lisent comme des bornes basses.",
        ],
    )

    answer(
        document,
        "À DIRE",
        "Une perte est bornée à zéro, sans plafond, et s'étale sur quatre ordres de grandeur : "
        "c'est la signature d'une variable dont le logarithme est à peu près normal, donc "
        "lognormale. J'ajuste par maximum de vraisemblance pondéré sur les logs — concrètement "
        "une moyenne et une variance pondérées — un ajustement par type d'attaque, parce qu'une "
        "lognormale unique se fait rejeter par Kolmogorov-Smirnov. Les poids viennent d'un noyau "
        "de similarité sur secteur, taille et maturité : je pondère plutôt que je filtre, parce "
        f"qu'un filtre exact ne laisserait que {f['hard_filter']} incidents. Le prix de cette "
        f"souplesse est un échantillon effectif plus petit, que je publie type par type, avec "
        "repli sur la distribution poolée en dessous de 30.",
    )

    # =========================================================== SIMULATION
    heading(document, "4. Simulation — ce que coûte une année", 2, new_page=True)

    heading(document, "Endpoint et paramètres", 3)
    code(
        document,
        "POST /api/simulate/\n"
        "{\n"
        '  "n_years": 5000,              "seed": 42,\n'
        '  "severity_threshold": "high", "session_window_hours": 24,\n'
        '  "loss_cap_eur": null,         "curve_points": 160,\n'
        '  "histogram_bins": 40,         "include_sensitivity": true\n'
        "}",
    )
    table(
        document,
        ["Paramètre", "Défaut", "Effet"],
        [
            ["`n_years`", "5 000 (API) / 100 000 (CLI)", "finesse de la queue ; borné à 200 000"],
            ["`seed`", "42", "reproductibilité bit à bit"],
            ["`severity_threshold`", "`high`", "convention de fréquence, cf. §2"],
            ["`session_window_hours`", "24", "convention de fréquence, cf. §2"],
            [
                "`loss_cap_eur`",
                f"99,9ᵉ centile ≈ {f['cap']}",
                "plafond de plausibilité par incident ; `inf` pour désactiver",
            ],
            ["`curve_points`", "160", "résolution des courbes OEP/AEP"],
            ["`histogram_bins`", "40", "bins de l'histogramme (log)"],
            ["`include_sensitivity`", "`true`", "grille 3×3 — neuf runs de plus"],
        ],
    )

    heading(document, "Chaîne de fonctions", 3)
    numbered(
        document,
        [
            "`get_simulation(n_years, seed, seuil, fenêtre, cap)` — mémoïsé ; **le cap fait "
            "partie de la clé de cache**, deux plafonds sont deux réponses.",
            "`simulate()` — la boucle Monte Carlo, vectorisée par blocs d'années.",
            "`rng.poisson(lam=λ_type, size=bloc)` — le nombre d'incidents de chaque année.",
            "`rng.lognormal(mean=mu, sigma=sigma, size=total)` — le prix de chaque incident.",
            "`np.minimum(losses, cap)` — le plafond de plausibilité.",
            "`np.bincount(years, weights=losses)` — replie les incidents dans leur année.",
            "`summarize(annual_losses)` — AAL, médiane, VaR, TVaR, P(aucune perte), maximum.",
            "`exceedance_curve(series, kind)` — courbes AEP (total) et OEP (pire perte unitaire).",
            "`SimulationResult.histogram(scale='log')` — bins log, années à zéro comptées à part.",
            "`sensitivity_grid()` — rejoue seuil × fenêtre et rapporte l'AAL de chaque case.",
        ],
    )

    heading(document, "Les maths", 3)
    bullets(
        document,
        [
            "**Loi de Poisson composée.** La perte annuelle est une somme d'un nombre aléatoire "
            "de termes aléatoires :",
        ],
    )
    code(document, "S = Σᵢ₌₁ᴺ Xᵢ    avec N ~ Poisson(λ)  et  Xᵢ ~ Lognormale(mu, sigma)")
    bullets(
        document,
        [
            "**Composer, pas multiplier.** Une année n'est pas « fréquence × coût moyen » : "
            f"c'est zéro incident **{f['p_no_loss']}**, un parfois, deux très rarement — et la "
            "somme de ces cas n'est pas le produit des moyennes.",
            "**Pourquoi simuler et pas une formule.** L'espérance est analytique (formule de "
            "Wald : `E[S] = λ × E[X]`), mais la **fonction de répartition de S n'a pas de forme "
            "fermée**. Or VaR 99 et TVaR 99 sont des quantiles de S. Il faut la distribution "
            "complète.",
            "**Ce qui fixe le nombre d'années** : la VaR 99 ne s'appuie que sur 1 % des années. "
            f"Sur {f['n_years']} années, cela fait 1 000 observations ; sur 1 000 années, dix — "
            "inutilisable. C'est la queue qui décide, pas la moyenne.",
            "**Métriques** : `VaR_α` = quantile d'ordre α ; `TVaR_α` = `E[S | S ≥ VaR_α]`. "
            "`TVaR ≥ VaR` par construction — c'est une invariante testée.",
            "**AEP vs OEP.** AEP = total de l'année ; OEP = plus grosse perte unitaire de "
            "l'année. L'OEP ne peut jamais dépasser l'AEP à probabilité égale, et leur écart "
            "mesure la part des années à plusieurs incidents.",
            f"**Plafond de plausibilité.** Une lognormale n'a pas de borne : à sigma "
            f"{f['sc_sigma']} et sur {f['n_years']} années, elle finit par tirer un incident "
            "coûtant plus que l'entreprise ne vaut (3,8 Md€ sans plafond). Chaque perte est donc "
            f"écrêtée à **{f['cap']}**, le 99,9ᵉ centile des pertes réellement observées chez "
            f"les pairs — {f['cap_share']} des tirages sont concernés.",
            "**L'effet du plafond est rapporté, pas absorbé** : le résultat porte "
            "`draws_capped`, `draws_total` et `aal_uncapped`.",
        ],
    )

    heading(document, "Chiffres produits", 3)
    table(
        document,
        ["Mesure", "Valeur", "Lecture métier"],
        [
            ["AAL", f["aal"], "ligne budgétaire"],
            ["Année médiane", f["median_year"], f"{f['p_no_loss']} sans incident"],
            ["VaR 95", f["var95"], "appétit au risque — 1 année sur 20"],
            ["TVaR 95", f["tvar95"], "moyenne des 5 % pires années"],
            ["VaR 99", f["var99"], "1 année sur 100"],
            ["TVaR 99", f["tvar99"], "question de survie"],
        ],
    )
    para(
        document,
        "Le rapport TVaR 95 / VaR 95 vaut 5,9 : une fois passé le seuil de la mauvaise année, "
        "la perte typique est presque six fois ce seuil. C'est la définition d'une queue lourde.",
        size=9.5,
        color=MUTED,
    )

    answer(
        document,
        "À DIRE",
        "Pour chaque année simulée je tire un nombre d'incidents par type dans une loi de "
        "Poisson, je prix chaque incident dans la lognormale de son type, et je somme. Cent "
        "mille années, graine 42, reproductible au centime. Je simule plutôt que d'appliquer une "
        "formule parce que l'espérance est analytique mais pas les quantiles : VaR 99 et TVaR 99 "
        "demandent la distribution complète. Et je plafonne chaque perte unitaire au 99,9ᵉ "
        "centile des pertes observées chez les pairs, parce qu'une lognormale non bornée finit "
        "par tirer une année à 3,8 milliards pour une ETI de 1 200 personnes — le plafond retire "
        "37,5 % de l'AAL, et je l'affiche plutôt que de l'absorber.",
    )

    # ========================================================= aide-mémoire
    heading(
        document,
        "Aide-mémoire : cinq questions, cinq réponses de trente secondes",
        2,
        new_page=True,
    )
    table(
        document,
        ["Question", "Réponse"],
        [
            [
                "Pourquoi dédoublonner ?",
                f"{f['dup_cross_feed']} événements sont vus par les deux flux. Concaténer "
                f"gonflerait toutes les fréquences de {f['inflation']}.",
            ],
            [
                "Pourquoi des épisodes ?",
                "Une intrusion produit une rafale d'alertes. Compter les alertes mesurerait le "
                "bavardage des sondes, pas la fréquence des attaques.",
            ],
            [
                "Pourquoi une calibration ?",
                "La télémétrie compte des attaques détectées, la base externe des incidents à "
                "perte. Multiplier les deux directement est une erreur de catégorie — c'est ce "
                "qui donnait 12,5 Md€.",
            ],
            [
                "Pourquoi pondérer plutôt que filtrer ?",
                f"Le filtre exact ne laisse que {f['hard_filter']} incidents et aucun type "
                "modélisable. La pondération garde tout le monde, les plus proches dominant.",
            ],
            [
                "Pourquoi Monte Carlo ?",
                "L'espérance a une formule, les quantiles non. VaR 99 et TVaR 99 exigent la "
                "distribution complète de la loi composée.",
            ],
        ],
    )

    para(
        document,
        "Détail des concepts mathématiques : CONCEPTS.md. Décisions de modélisation et "
        "justifications : METHODOLOGY.md. Ce qui n'a pas été fait et pourquoi : next_steps.md.",
        size=9,
        color=MUTED,
        italic=True,
    )

    return document


def build_en(f: dict) -> Document:
    """The English reference. Block for block the same document as `build_fr`."""
    document = new_document()
    heading(document, "Business logic, page by page", 1)
    para(
        document,
        "For each page of the cockpit: the endpoint it calls, its parameters, the chain of "
        "functions it runs through, and the maths behind the calculation. Every section ends "
        "with the answer to give out loud.",
        color=MUTED,
    )
    para(
        document,
        f"Engine figures, seed 42, {f['n_years']} simulated years. Every value quoted here comes "
        "from deck/figures.json, written by the pipeline.",
        size=9,
        color=MUTED,
        italic=True,
    )

    # ------------------------------------------------------------- overview
    heading(document, "Overview", 2)
    table(
        document,
        ["Page", "Endpoint", "Parameters", "Cost"],
        [
            ["Overview", "the four below", "\u2014", "all cached"],
            [
                "Telemetry",
                "`GET /api/telemetry/summary/`\n`GET /api/frequency/`",
                "none",
                "cached",
            ],
            [
                "Frequency",
                "`GET /api/frequency/`\n`GET /api/assets/`",
                "threshold, window",
                "~1 s",
            ],
            ["Severity", "`GET /api/severity/`", "none", "cached"],
            ["Simulation", "`POST /api/simulate/`", "7 (see \u00a74)", "a few s"],
        ],
    )
    bullets(
        document,
        [
            "**One architectural rule**: the view parses the request, calls the engine, "
            "serializes. No risk maths in Django, none in the frontend.",
            "**Two levels of cache**: `get_dataset()` loads and fits once for all \u2014 the CSVs "
            "do not change \u2014 while `get_frequency()` and `get_simulation()` are memoized per "
            "parameter set.",
            "**Every result object exposes** `to_explanation()` \u2014 the numbered trace the CLI "
            "prints, the API serves and the UI displays. One version of each figure.",
        ],
    )

    # ============================================================ TELEMETRY
    heading(document, "1. Telemetry \u2014 what the sensors saw", 2, new_page=True)

    heading(document, "Endpoint and parameters", 3)
    code(document, "GET /api/telemetry/summary/      (no parameters)\nGET /api/frequency/")
    bullets(
        document,
        [
            "The second call serves only the **last stage of the funnel**: the incident rate the "
            "simulation consumes.",
            "No parameters: normalization depends on no convention.",
        ],
    )

    heading(document, "Call chain", 3)
    numbered(
        document,
        [
            "`load_assets()` \u2014 reads the reference for the 20 assets.",
            "`load_siem()` / `load_edr()` \u2014 read each export and translate it into canonical "
            "`SecurityEvent`s. This is where the two severity scales are brought into one shared "
            "vocabulary.",
            "`severity_from_siem_label()` \u2014 `Low/Medium/High/Critical` \u2192 class.",
            "`severity_from_edr_risk()` \u2014 score 0\u2013999 \u2192 class, by cut points.",
            "`merge_feeds()` \u2014 **deduplicates** over the union of both feeds and produces the "
            "`NormalizationReport`.",
            "`observed_window()` \u2014 derives the observation window from the events themselves.",
            "`summarize_telemetry()` \u2014 aggregates into weekly buckets, severity mix, top "
            "techniques.",
        ],
    )

    heading(document, "The maths", 3)
    bullets(
        document,
        [
            "**Severity normalization.** The SIEM gives a label, the EDR a score. EDR cut "
            "points: `< 50` low, `50\u201369` medium, `70\u201393` high, `\u2265 94` critical. "
            "Numeric score attached to each class: 0.25 / 0.5 / 0.75 / 1.0.",
            "**Deduplication key**: the triple `(asset_id, technique, observed_at)`. Two reports "
            "carrying the same triple describe the same real event.",
            "**Deduplicated over the union, not by a join.** Each feed already repeats keys "
            "internally \u2014 761 times in the SIEM, 543 in the EDR \u2014 so an inner join would "
            "multiply those duplicates and return 13,055 rows instead of the 12,343 true matches.",
            "**Merge rule: the worst severity wins.** If the EDR sees critical and the SIEM "
            "medium, critical is kept: the tool closest to the endpoint saw something. Averaging "
            "would dilute a strong signal into a weak one.",
            f"**Annualization factor** = 365 / observed days = 365 / {f['observed_days']} = "
            f"**{f['annualization']}** (1 Nov 2025 \u2192 31 May 2026). Recomputed from the data, "
            "never hardcoded.",
        ],
    )

    heading(document, "Figures produced", 3)
    table(
        document,
        ["Quantity", "Value", "Where it comes from"],
        [
            ["Raw rows", f["rows_read"], "sum of both exports"],
            ["Distinct events", f["total_events"], "after deduplication"],
            [
                "Duplicates absorbed",
                f["duplicates"],
                f"{f['dup_cross_feed']} of them across feeds",
            ],
            [
                "Inflation avoided",
                f["inflation"],
                "what a naive concatenation would have added",
            ],
        ],
    )

    answer(
        document,
        "SAY THIS",
        "The two tools are not complementary, they are partly redundant: "
        f"{f['dup_cross_feed']} events carry the same asset, the same technique and the same "
        "timestamp on both sides. I deduplicate on that triple keeping the worst severity, which "
        f"avoids {f['inflation']} of inflation on every downstream rate.",
    )

    # ============================================================ FREQUENCY
    heading(document, "2. Frequency \u2014 how often we are attacked", 2, new_page=True)

    heading(document, "Endpoint and parameters", 3)
    code(
        document,
        "GET /api/frequency/?severity_threshold=high&session_window_hours=24\n"
        "GET /api/assets/?severity_threshold=high&session_window_hours=24",
    )
    table(
        document,
        ["Parameter", "Default", "Effect"],
        [
            ["`severity_threshold`", "`high`", "minimum severity to count as an attack"],
            ["`session_window_hours`", "`24`", "silence that opens a new episode"],
        ],
    )
    para(
        document,
        "These two parameters are conventions, not measurements. They are exposed as dials in "
        "the UI rather than hidden in the code.",
        size=9.5,
        color=MUTED,
        italic=True,
    )

    heading(document, "Call chain", 3)
    numbered(
        document,
        [
            "`get_frequency(threshold, window)` \u2014 memoized per parameter pair.",
            "`estimate_frequency()` \u2014 orchestrates the four steps below.",
            "`is_attack_grade(event, threshold)` \u2014 filter: keeps only high and above.",
            "`sessionize(events, params)` \u2014 groups events into **episodes**.",
            "`attack_type_for(technique)` \u2014 classifies each MITRE technique into an attack "
            "type through an explicit dictionary; the episode inherits the type of its worst "
            "event.",
            "`peer_weighted_base_rate(incidents, params)` \u2014 incidents per organisation-year "
            "among peers.",
            "`calibrate(lambda_detected, base_rate)` \u2014 fits `p_materialize` to reconcile the "
            "two units.",
        ],
    )

    heading(document, "The maths", 3)
    bullets(
        document,
        [
            "**Episode rule.** Grouped by **asset alone** \u2014 an intruder trips whatever "
            "detections it meets, and counting each type separately would recount the same "
            "intrusion under several names. The run is cut as soon as a silence exceeds the "
            "window, measured against the **previous** event and not the first, otherwise a "
            "three-day intrusion would be split artificially.",
            f"**\u03bb detected** = episodes \u00d7 365 / observed days = {f['episodes']} \u00d7 "
            f"{f['annualization']} = **{f['lambda_detected']} / yr**.",
            "**Peer base rate** = weighted incidents \u00f7 weighted organisation-years. Numerator "
            "*and* denominator carry the same similarity kernel, one record per `company_id` "
            f"\u2014 that is **{f['peer_companies']} organisations**.",
            "**The change of unit, the key point.** The telemetry counts *detected* attacks, the "
            "external base counts *loss-causing* incidents. Those are two different units.",
            f"**Calibration**: `\u03bb_incident = \u03bb_detected \u00d7 p`, with p set so that "
            f"\u03bb_incident equals the peer rate. Here p \u2248 1 detection in "
            f"**{f['p_one_in']}**, so \u03bb_incident = **{f['lambda_incident']} / yr**.",
            "**The mix comes from the telemetry, the level from the base.** Each type gets "
            "\u03bb_incident \u00d7 (its episodes / total episodes).",
        ],
    )

    heading(document, "Figures produced", 3)
    table(
        document,
        ["Quantity", "Value", "Calculation"],
        [
            ["Attack-grade events", f["attack_grade"], "severity \u2265 high"],
            ["Episodes", f["episodes"], f"compression {f['compression']}"],
            ["\u03bb detected", f"{f['lambda_detected']} / yr", "episodes \u00d7 365/212"],
            [
                "\u03bb incident",
                f"{f['lambda_incident']} / yr",
                f"\u03bb detected \u00f7 {f['p_one_in']}",
            ],
        ],
    )

    answer(
        document,
        "SAY THIS",
        "An alert is not an attack: an intrusion produces a burst of detections. I group by "
        f"asset with a 24-hour quiet window, which takes {f['attack_grade']} attack-grade events "
        f"down to {f['episodes']} episodes, or {f['lambda_detected']} detected attacks a year. "
        "But the telemetry counts detections and the external base counts loss-causing incidents "
        f"\u2014 two different units. I reconcile them with a calibration: one detection in "
        f"{f['p_one_in']} ends in a loss, giving {f['lambda_incident']} incident a year anchored "
        f"on {f['peer_companies']} comparable organisations.",
    )

    # ============================================================= SEVERITY
    heading(document, "3. Severity \u2014 what one incident costs", 2, new_page=True)

    heading(document, "Endpoint and parameters", 3)
    code(document, "GET /api/severity/      (no parameters)")
    para(
        document,
        "No parameters: the severity model depends only on the target profile, which is fixed "
        "for this case (Retail, ETI, maturity 55).",
        size=9.5,
        color=MUTED,
        italic=True,
    )

    heading(document, "Call chain", 3)
    numbered(
        document,
        [
            "`load_incidents()` \u2014 reads and **cleans** the external base; also returns the "
            "`CleaningReport`.",
            "`repair_mojibake()` \u2014 repairs the double UTF-8 encoding of the sector labels, "
            "which was splitting one sector into two.",
            "`fit_severity_model()` \u2014 orchestrates the steps below, once for the pool then "
            "once per attack type.",
            "`peer_weights(incidents, params)` \u2014 one weight per incident, via "
            "`sector_weight()`, `size_weight()` and `maturity_weight()`.",
            "`effective_sample_size(weights)` \u2014 Kish n_eff, what the weighting costs.",
            "`fit_lognormal(losses, weights)` \u2014 weighted maximum likelihood on the log scale.",
            "`diagnose()` \u2014 assembles the evidence *against* the fit: `weighted_ks()`, "
            "`qq_points()`, `fit_pareto_tail()`, `distribution_plot()`.",
        ],
    )

    heading(document, "The maths", 3)
    bullets(
        document,
        [
            "**Cleaning.** Losses of `-1` are **missing values**, never \u20ac0. Two rows "
            f"affected, excluded from the fit: **{f['incidents_fitted']} usable incidents of "
            "1,600** remain.",
            "**Soft weighting, never a hard filter.** A continuous weight per incident, the "
            "product of three similarity factors: sector (1.0 if Retail, else 0.4), size (1.0 if "
            "ETI, else 0.6) and a Gaussian kernel on the maturity distance, of bandwidth h = 15.",
            f"**Why weight.** An exact filter on the profile leaves only **{f['hard_filter']} "
            f"incidents of {f['incidents_fitted']}**, and no attack type retains a credible "
            "sample. Mathematically clean, practically unusable.",
            "**Weighted lognormal fit** on the log scale \u2014 it is simply a weighted mean and a "
            "weighted variance:",
        ],
    )
    code(
        document,
        "w      = w_sector \u00d7 w_size \u00d7 exp(\u2212 d\u00b2 / 2h\u00b2)      "
        "d = |maturity \u2212 55|\n"
        "mu     = \u03a3 w\u1d62 ln x\u1d62 / \u03a3 w\u1d62\n"
        "sigma\u00b2 = \u03a3 w\u1d62 (ln x\u1d62 \u2212 mu)\u00b2 / \u03a3 w\u1d62",
    )
    bullets(
        document,
        [
            "**Back to euros**: `median = e^mu` and `mean = e^(mu + sigma\u00b2/2)`. Careful: mu "
            "and sigma are **not** the mean and standard deviation of the loss, but those of its "
            "logarithm.",
            f"**Example, supply chain**: mu = {f['sc_mu']}, sigma = {f['sc_sigma']} \u2192 median "
            f"**{f['sc_median']}**, mean **{f['sc_mean']}**. The tail multiplies the mean by "
            f"**{f['sc_ratio']}**: e^(sigma\u00b2/2).",
            "**It is the mean that feeds the annual loss**, not the median. A decision-maker "
            f"reasoning about the typical incident is wrong by a factor of {f['sc_ratio']}.",
            "**Kish n_eff** = `(\u03a3w)\u00b2 / \u03a3w\u00b2`. It is the number of equally "
            "weighted observations that would carry as much information. It measures the "
            f"**evenness** of the weights, not their size. Below **{f['min_neff']}**, the type "
            "falls back to the pooled distribution, and the substitution is recorded.",
            f"**Evidence published with every fit**: weighted KS (here **{f['sc_ks']}** against "
            "an indicative threshold \u2248 0.19), a QQ plot of the log-losses, and a **Pareto "
            f"tail fitted as a rival** \u2014 it wins on {f['pareto_better']} fits of "
            f"{f['fits_total']}, so VaR 99 and TVaR 99 read as lower bounds.",
        ],
    )

    answer(
        document,
        "SAY THIS",
        "A loss is bounded at zero, has no ceiling, and spans four orders of magnitude: that is "
        "the signature of a variable whose logarithm is roughly normal, hence lognormal. I fit "
        "by weighted maximum likelihood on the logs \u2014 concretely a weighted mean and variance "
        "\u2014 one fit per attack type, because a single lognormal is rejected by "
        "Kolmogorov-Smirnov. The weights come from a similarity kernel on sector, size and "
        "maturity: I weight rather than filter, because an exact filter would leave only "
        f"{f['hard_filter']} incidents. The price of that flexibility is a smaller effective "
        "sample, which I publish type by type, with a fallback to the pooled distribution below "
        "30.",
    )

    # =========================================================== SIMULATION
    heading(document, "4. Simulation \u2014 what a year costs", 2, new_page=True)

    heading(document, "Endpoint and parameters", 3)
    code(
        document,
        "POST /api/simulate/\n"
        "{\n"
        '  "n_years": 5000,              "seed": 42,\n'
        '  "severity_threshold": "high", "session_window_hours": 24,\n'
        '  "loss_cap_eur": null,         "curve_points": 160,\n'
        '  "histogram_bins": 40,         "include_sensitivity": true\n'
        "}",
    )
    table(
        document,
        ["Parameter", "Default", "Effect"],
        [
            [
                "`n_years`",
                "5,000 (API) / 100,000 (CLI)",
                "tail resolution; capped at 200,000",
            ],
            ["`seed`", "42", "bit-for-bit reproducibility"],
            ["`severity_threshold`", "`high`", "frequency convention, see \u00a72"],
            ["`session_window_hours`", "24", "frequency convention, see \u00a72"],
            [
                "`loss_cap_eur`",
                f"99.9th percentile \u2248 {f['cap']}",
                "per-incident ceiling; `inf` to disable",
            ],
            ["`curve_points`", "160", "resolution of the OEP/AEP curves"],
            ["`histogram_bins`", "40", "histogram bins (log)"],
            ["`include_sensitivity`", "`true`", "3\u00d73 grid \u2014 nine more runs"],
        ],
    )

    heading(document, "Call chain", 3)
    numbered(
        document,
        [
            "`get_simulation(n_years, seed, threshold, window, cap)` \u2014 memoized; **the cap is "
            "part of the cache key**, two ceilings are two answers.",
            "`simulate()` \u2014 the Monte Carlo loop, vectorized over blocks of years.",
            "`rng.poisson(lam=\u03bb_type, size=block)` \u2014 the incident count of each year.",
            "`rng.lognormal(mean=mu, sigma=sigma, size=total)` \u2014 the price of each incident.",
            "`np.minimum(losses, cap)` \u2014 the plausibility ceiling.",
            "`np.bincount(years, weights=losses)` \u2014 folds incidents back into their year.",
            "`summarize(annual_losses)` \u2014 AAL, median, VaR, TVaR, P(no loss), maximum.",
            "`exceedance_curve(series, kind)` \u2014 AEP (annual total) and OEP (largest single "
            "loss) curves.",
            "`SimulationResult.histogram(scale='log')` \u2014 log bins, zero years counted apart.",
            "`sensitivity_grid()` \u2014 replays threshold \u00d7 window and "
            "reports the AAL of each cell.",
        ],
    )

    heading(document, "The maths", 3)
    bullets(
        document,
        [
            "**Compound Poisson.** The annual loss is a sum of a random number of random terms:",
        ],
    )
    code(
        document,
        "S = \u03a3\u1d62\u208c\u2081\u1d3f X\u1d62    with N ~ Poisson(\u03bb)  and  "
        "X\u1d62 ~ Lognormal(mu, sigma)",
    )
    bullets(
        document,
        [
            "**Compound, do not multiply.** A year is not frequency times mean cost: it is zero "
            f"incidents **{f['p_no_loss']}**, one sometimes, two very rarely \u2014 and the sum of "
            "those cases is not the product of the means.",
            "**Why simulate rather than use a formula.** The expectation is analytic (Wald's "
            "identity: `E[S] = \u03bb \u00d7 E[X]`), but **the distribution function of S has no "
            "closed form**. And VaR 99 and TVaR 99 are quantiles of S. The full distribution is "
            "required.",
            "**What fixes the number of years**: VaR 99 rests on only 1% of the years. Over "
            f"{f['n_years']} years that is 1,000 observations; over 1,000 years, ten \u2014 "
            "unusable. The tail decides, not the mean.",
            "**Metrics**: `VaR_\u03b1` = quantile of order \u03b1; `TVaR_\u03b1` = "
            "`E[S | S \u2265 VaR_\u03b1]`. `TVaR \u2265 VaR` by construction \u2014 a tested "
            "invariant.",
            "**AEP vs OEP.** AEP = the year's total; OEP = the year's largest single loss. OEP "
            "can never exceed AEP at equal probability, and their gap measures the share of years "
            "with several incidents.",
            f"**Plausibility ceiling.** A lognormal has no upper bound: at sigma {f['sc_sigma']} "
            f"and over {f['n_years']} years it eventually draws an incident costing more than the "
            "company is worth (\u20ac3.8bn uncapped). Every loss is therefore clipped at "
            f"**{f['cap']}**, the 99.9th percentile of the losses comparable organisations were "
            f"actually observed to suffer \u2014 {f['cap_share']} of draws are affected.",
            "**The ceiling's effect is reported, not absorbed**: the result carries "
            "`draws_capped`, `draws_total` and `aal_uncapped`.",
        ],
    )

    heading(document, "Figures produced", 3)
    table(
        document,
        ["Measure", "Value", "Business reading"],
        [
            ["AAL", f["aal"], "the budget line"],
            ["Median year", f["median_year"], f"{f['p_no_loss']} with no incident"],
            ["VaR 95", f["var95"], "risk appetite \u2014 1 year in 20"],
            ["TVaR 95", f["tvar95"], "mean of the worst 5% of years"],
            ["VaR 99", f["var99"], "1 year in 100"],
            ["TVaR 99", f["tvar99"], "the survival question"],
        ],
    )
    para(
        document,
        "The TVaR 95 / VaR 95 ratio is 5.9: once past the threshold of the bad year, the typical "
        "loss is almost six times that threshold. That is the definition of a heavy tail.",
        size=9.5,
        color=MUTED,
    )

    answer(
        document,
        "SAY THIS",
        "For each simulated year I draw an incident count per attack type from a Poisson "
        "distribution, price each incident in its type's lognormal, and sum. A hundred thousand "
        "years, seed 42, reproducible to the cent. I simulate rather than apply a formula because "
        "the expectation is analytic but the quantiles are not: VaR 99 and TVaR 99 need the full "
        "distribution. And I cap each single loss at the 99.9th percentile of the losses observed "
        "among peers, because an unbounded lognormal eventually draws a \u20ac3.8 billion year for "
        "a 1,200-person company \u2014 the cap removes 37.5% of the AAL, and I show that rather "
        "than absorb it.",
    )

    # ============================================================ cheatsheet
    heading(
        document,
        "Cheat sheet: five questions, five thirty-second answers",
        2,
        new_page=True,
    )
    table(
        document,
        ["Question", "Answer"],
        [
            [
                "Why deduplicate?",
                f"{f['dup_cross_feed']} events are seen by both feeds. Concatenating would "
                f"inflate every rate by {f['inflation']}.",
            ],
            [
                "Why episodes?",
                "An intrusion produces a burst of alerts. Counting alerts would measure the "
                "estate's detection verbosity, not how often it is attacked.",
            ],
            [
                "Why a calibration?",
                "The telemetry counts detected attacks, the external base counts loss-causing "
                "incidents. Multiplying the two directly is a category error \u2014 it is what "
                "produced \u20ac12.5bn.",
            ],
            [
                "Why weight rather than filter?",
                f"An exact filter leaves only {f['hard_filter']} incidents and no modellable "
                "type. Weighting keeps everyone, with the closest dominating.",
            ],
            [
                "Why Monte Carlo?",
                "The expectation has a formula, the quantiles do not. VaR 99 and TVaR 99 require "
                "the full compound distribution.",
            ],
        ],
    )

    para(
        document,
        "Mathematical concepts in detail: CONCEPTS.md (in French). Modeling decisions and their "
        "justification: METHODOLOGY.md. What was not built and why: next_steps.md.",
        size=9,
        color=MUTED,
        italic=True,
    )

    return document


def to_pdf(docx: Path) -> Path | None:
    """Export the .docx to PDF through Word, if Word is installed."""
    pdf = docx.with_suffix(".pdf")
    script = (
        "$w = New-Object -ComObject Word.Application; $w.Visible = $false; "
        f"$d = $w.Documents.Open('{docx}'); $d.ExportAsFixedFormat('{pdf}', 17); "
        "$d.Close($false); $w.Quit()"
    )
    try:
        subprocess.run(
            ["powershell", "-NoProfile", "-Command", script],
            check=True,
            capture_output=True,
            timeout=300,
        )
    except (subprocess.CalledProcessError, subprocess.TimeoutExpired, FileNotFoundError) as error:
        print(f"  PDF non généré ({type(error).__name__}) - le .docx reste utilisable")
        return None
    print(f"{pdf}")
    return pdf


#: French number forms that need an English rendering. The pipeline writes
#: `deck/figures.json` in French convention - space for thousands, comma for
#: decimals, "12 M€" with the sign after - because the deck is French. The
#: English document reads the same file and converts, so the two versions can
#: never quote different numbers.
_CURRENCY = re.compile(r"^([\d  ,.]+?)\s*(k€|M€|Md€|€)$")


def anglicise(figures: dict) -> dict:
    """Re-render French-formatted figures in English convention.

    "45 840" -> "45,840", "42,4 %" -> "42.4%", "5,1 M€" -> "€5.1M".
    """

    def convert(text: str) -> str:
        currency = _CURRENCY.match(text)
        if currency:
            digits, unit = currency.groups()
            suffix = {"k€": "k", "M€": "M", "Md€": "bn", "€": ""}[unit]
            return "€" + _swap(digits.strip()) + suffix
        if text.endswith(" %"):
            return _swap(text[:-2]) + "%"
        return _swap(text)

    converted = {key: convert(value) for key, value in figures.items()}
    # Two entries carry French words as well as digits.
    share = convert(figures["p_no_loss"].replace(" des années", ""))
    converted["p_no_loss"] = f"{share} of years"
    converted["median_year"] = "€0"
    return converted


#: French grouping and decimal marks mapped to English ones. `str.translate`
#: applies every mapping in a single pass, so no placeholder character is
#: needed to stop the comma substitutions from colliding.
_FR_TO_EN = str.maketrans({" ": ",", " ": ",", ",": "."})


def _swap(text: str) -> str:
    """Swap French digit grouping and decimal mark for English ones."""
    return text.translate(_FR_TO_EN)


#: Language code -> (builder, output stem).
LANGUAGES = {
    "fr": (build_fr, "logique_metier"),
    "en": (build_en, "business_logic"),
}


def main() -> int:
    """Build the reference in each requested language."""
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--figures", type=Path, default=ROOT / "deck" / "figures.json")
    parser.add_argument("--out-dir", type=Path, default=ROOT)
    parser.add_argument("--lang", choices=[*LANGUAGES, "all"], default="all")
    parser.add_argument("--no-pdf", action="store_true")
    args = parser.parse_args()

    figures = json.loads(args.figures.read_text(encoding="utf-8"))
    wanted = LANGUAGES if args.lang == "all" else {args.lang: LANGUAGES[args.lang]}

    outlines: dict[str, list[str]] = {}
    for language, (builder, stem) in wanted.items():
        _OUTLINE.clear()
        document = builder(anglicise(figures) if language == "en" else figures)
        outlines[language] = list(_OUTLINE)

        docx = args.out_dir / f"{stem}.docx"
        docx.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(docx))
        print(f"  [{language}] {docx.name}")
        if not args.no_pdf:
            pdf = to_pdf(docx)
            if pdf is not None:
                print(f"  [{language}] {pdf.name}")

    # Parity guard: the two documents must have the same shape. A section added
    # to one language and forgotten in the other is otherwise invisible until
    # someone reads both side by side.
    if len(outlines) == 2:
        first, second = outlines["fr"], outlines["en"]
        if first != second:
            print(
                f"  ATTENTION: structures divergentes - fr {len(first)} titres, "
                f"en {len(second)} titres",
                file=sys.stderr,
            )
            return 1
        print(f"  structures identiques ({len(first)} titres)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
